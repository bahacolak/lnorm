"""
docx_writer.py - Word output generation.
"""

from __future__ import annotations

import logging
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

from .articles_parser import Article
from .extractor import BoardMember, CompanyInfo

logger = logging.getLogger(__name__)

FONT_NAME = "Calibri"
TITLE_SIZE = Pt(16)
HEADING_SIZE = Pt(13)
BODY_SIZE = Pt(11)
SMALL_SIZE = Pt(9)
COLOR_PRIMARY = RGBColor(0x1A, 0x3C, 0x6E)
COLOR_WARNING = RGBColor(0xCC, 0x00, 0x00)
COLOR_CHANGED = RGBColor(0xE6, 0x7E, 0x22)
COLOR_SOURCE = RGBColor(0x66, 0x66, 0x66)


def _apply_base_style(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = BODY_SIZE


def _add_title(doc: Document, title: str) -> None:
    heading = doc.add_heading(title, level=1)
    for run in heading.runs:
        run.font.color.rgb = COLOR_PRIMARY
        run.font.size = TITLE_SIZE


def _format_cell(cell, text: str, *, bold: bool = False, color=None, size=None) -> None:
    cell.text = ""
    run = cell.paragraphs[0].add_run(text)
    run.font.name = FONT_NAME
    run.font.size = size or BODY_SIZE
    run.bold = bold
    if color:
        run.font.color.rgb = color


def write_sirket_bilgileri(info: CompanyInfo, output_path: str = "output/sirket_bilgileri.docx") -> Path:
    doc = Document()
    _apply_base_style(doc)
    _add_title(doc, "Güncel Şirket Bilgileri")

    rows = [
        ("Ticaret Unvanı", info.ticaret_unvani),
        ("Şirket Türü", info.sirket_turu),
        ("MERSİS Numarası", info.mersis_no),
        ("Ticaret Sicil Müdürlüğü", info.ticaret_sicil_mudurlugu),
        ("Ticaret Sicil Numarası", info.ticaret_sicil_no),
        ("Adres", info.adres),
        ("Mevcut Sermaye", info.sermaye),
        ("Kuruluş Tarihi", info.kurulus_tarihi),
        ("Denetçi", info.denetci or "—"),
    ]

    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for idx, (label, value) in enumerate(rows):
        _format_cell(table.rows[idx].cells[0], label, bold=True)
        if value and "[DOĞRULANAMADI]" in str(value):
            _format_cell(table.rows[idx].cells[1], str(value), color=COLOR_WARNING)
        else:
            _format_cell(table.rows[idx].cells[1], str(value or "—"))

    doc.add_paragraph()
    doc.add_paragraph("Alan Bazlı Kaynaklar")
    for field_name, source in sorted(info.field_sources.items()):
        para = doc.add_paragraph(style=None)
        run = para.add_run(f"{field_name}: {source.get('ttsg_date', '—')} — {source.get('pdf', '—')}")
        run.font.size = SMALL_SIZE
        run.font.color.rgb = COLOR_SOURCE

    return _save(doc, output_path)


def write_yonetim_kurulu(members: list[BoardMember], output_path: str = "output/yonetim_kurulu.docx") -> Path:
    doc = Document()
    _apply_base_style(doc)
    _add_title(doc, "Yönetim Kurulu Üyeleri")

    if not members:
        doc.add_paragraph("Yönetim kurulu üyesi bilgisi bulunamadı. Manuel doğrulama gerekli.")
        return _save(doc, output_path)

    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = [
        "Ad Soyad / Unvan",
        "Görevi / Unvanı",
        "Görev Bitiş Tarihi",
        "Atandığı TTSG Tarih / Sayı",
        "Kaynak PDF Bağlantısı",
    ]
    for idx, header in enumerate(headers):
        _format_cell(table.rows[0].cells[idx], header, bold=True, color=COLOR_PRIMARY)

    for member in members:
        row = table.add_row()
        display_name = member.name
        if member.entity_type == "legal_entity" and member.representative:
            display_name = f"{member.name} (adına hareket edecek gerçek kişi: {member.representative})"
        _format_cell(
            row.cells[0],
            display_name,
            color=COLOR_WARNING if "[DOĞRULANAMADI]" in display_name else None,
        )
        _format_cell(row.cells[1], member.role)
        _format_cell(row.cells[2], member.term_end or "—")
        appointment_ref = member.appointment_ttsg_date or "—"
        if member.appointment_ttsg_no:
            appointment_ref = f"{appointment_ref} / {member.appointment_ttsg_no}"
        _format_cell(row.cells[3], appointment_ref)
        row.cells[4].text = ""
        paragraph = row.cells[4].paragraphs[0]
        if member.source_pdf_link:
            _add_hyperlink(paragraph, member.kaynak_pdf, member.source_pdf_link)
        else:
            paragraph.add_run(member.kaynak_pdf or "—")

    return _save(doc, output_path)


def write_esas_sozlesme(articles: list[Article], sources: dict, output_path: str = "output/esas_sozlesme.docx") -> Path:
    doc = Document()
    _apply_base_style(doc)
    _add_title(doc, "Esas Sözleşme (Konsolide)")

    if len(articles) != 16:
        warning = doc.add_paragraph()
        run = warning.add_run(
            f"Uyarı: 16 madde bekleniyordu, {len(articles)} madde üretildi. Manuel doğrulama gerekli."
        )
        run.font.color.rgb = COLOR_WARNING

    if not articles:
        doc.add_paragraph("Esas sözleşme maddesi bulunamadı.")
        return _save(doc, output_path)

    for article in articles:
        source = sources.get(article.madde_no, {})
        heading_text = f"MADDE {article.madde_no}"
        if article.baslik:
            heading_text += f" — {article.baslik}"
        if source.get("degistirildi"):
            heading_text += " [DEĞİŞTİRİLDİ]"

        heading = doc.add_heading(heading_text, level=2)
        for run in heading.runs:
            run.font.size = HEADING_SIZE
            if source.get("degistirildi"):
                run.font.color.rgb = COLOR_CHANGED

        paragraph = doc.add_paragraph()
        run = paragraph.add_run(article.icerik)
        run.font.name = FONT_NAME
        run.font.size = BODY_SIZE

        source_para = doc.add_paragraph()
        source_text = f"Kaynak: TTSG [{source.get('tarih', article.kaynak_tarih or '—')}]"
        if source.get("sayi"):
            source_text += f" Sayı {source['sayi']}"
        if source.get("kaynak_pdf"):
            source_text += f" — {source['kaynak_pdf']}"
        source_run = source_para.add_run(source_text)
        source_run.font.size = SMALL_SIZE
        source_run.font.color.rgb = COLOR_SOURCE
        source_run.italic = True

    return _save(doc, output_path)


def _save(doc: Document, output_path: str) -> Path:
    path = Path(output_path)
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))
    logger.info("Kaydedildi: %s", path)
    return path


def _add_hyperlink(paragraph, text: str, url: str) -> None:
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    r_pr.append(color)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)
    new_run.append(r_pr)

    text_element = OxmlElement("w:t")
    text_element.text = text
    new_run.append(text_element)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
